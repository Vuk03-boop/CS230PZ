# -*- coding: utf-8 -*-
"""Generise seminarski rad CS230 kao .docx."""
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIG = os.path.join(ROOT, "figures")
OUT = os.path.join(ROOT, "../CS230-Seminarski-Vuk-Trifunovic-6223.docx")

FONT = "Times New Roman"

doc = Document()

# ---------------------------------------------------------------- stilovi ---
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(12)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
st.element.rPr.rFonts.set(qn("w:cs"), FONT)
pf = st.paragraph_format
pf.line_spacing = 1.0
pf.space_after = Pt(6)
pf.space_before = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, before, after in (("Heading 1", 15, 18, 8),
                                  ("Heading 2", 13, 12, 6)):
    s = doc.styles[name]
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    # Naslovni stilovi nasledjuju tematski font; w:*Theme ima prednost nad
    # w:ascii, pa se atributi teme moraju ukloniti da bi TNR bio primenjen.
    rf = s.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rf.attrib.pop(qn("w:" + a), None)
    for a in ("ascii", "hAnsi", "eastAsia", "cs"):
        rf.set(qn("w:" + a), FONT)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.line_spacing = 1.0
    s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    s.paragraph_format.keep_with_next = True

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin = sec.bottom_margin = Cm(2.5)
sec.different_first_page_header_footer = True

USABLE = Inches(6.3)


# ---------------------------------------------------------------- pomocne ---
def para(text="", style=None, align=None, size=None, bold=False, italic=False,
         space_after=None, space_before=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = FONT
        if size:
            r.font.size = Pt(size)
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def p(text):
    return para(text)


def formula(text):
    return para(text, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True,
                space_before=6, space_after=6)


def code(lines):
    for i, line in enumerate(lines):
        cp = doc.add_paragraph()
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.paragraph_format.space_after = Pt(10 if i == len(lines) - 1 else 0)
        cp.paragraph_format.space_before = Pt(6 if i == 0 else 0)
        cp.paragraph_format.left_indent = Cm(0.8)
        cp.paragraph_format.line_spacing = 1.0
        r = cp.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def bullets(items):
    for it in items:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(3)
        bp.paragraph_format.line_spacing = 1.0
        bp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = bp.add_run(it)
        r.font.name = FONT
        r.font.size = Pt(12)


def caption(text):
    return para(text, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True,
                space_before=2, space_after=12)


def figure(fname, cap, width_in=6.3):
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(8)
    fp.paragraph_format.space_after = Pt(2)
    fp.paragraph_format.keep_with_next = True
    fp.add_run().add_picture(os.path.join(FIG, fname), width=Inches(width_in))
    caption(cap)


def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def table(cap, headers, rows, widths, align_right=()):
    """widths: relativni udeli kolona; normalizuju se na sirinu stranice."""
    tc = para(cap, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, italic=True,
              space_before=8, space_after=4)
    tc.paragraph_format.keep_with_next = True
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    total = sum(widths)
    cw = [Inches(6.3 * w / total) for w in widths]

    def fill(cells, values, bold, fill_hex=None):
        for i, (c, v) in enumerate(zip(cells, values)):
            c.width = cw[i]
            cp = c.paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.line_spacing = 1.0
            if i in align_right:
                cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif bold:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = cp.add_run(str(v))
            r.bold = bold
            r.font.name = FONT
            r.font.size = Pt(10)
            if fill_hex:
                shade(c, fill_hex)

    fill(t.rows[0].cells, headers, True, "D9D9D9")
    for row in rows:
        fill(t.add_row().cells, row, False)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def field(paragraph, instr):
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar")
    fc.set(qn("w:fldCharType"), "begin")
    r._r.append(fc)
    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r2._r.append(it)
    r3 = paragraph.add_run()
    fs = OxmlElement("w:fldChar")
    fs.set(qn("w:fldCharType"), "separate")
    r3._r.append(fs)
    r4 = paragraph.add_run("1")
    r4.font.name = FONT
    r5 = paragraph.add_run()
    fe = OxmlElement("w:fldChar")
    fe.set(qn("w:fldCharType"), "end")
    r5._r.append(fe)


def pagebreak():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ------------------------------------------------------------ 1. NASLOVNA ---
for _ in range(3):
    para(space_after=0)
para("[Naziv univerziteta / fakulteta]", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("[Studijski program]", align=WD_ALIGN_PARAGRAPH.CENTER, size=13,
     space_after=90)

para("CS230 – Distribuirani sistemi", align=WD_ALIGN_PARAGRAPH.CENTER, size=14,
     bold=True, space_after=6)
para("SEMINARSKI RAD", align=WD_ALIGN_PARAGRAPH.CENTER, size=13,
     space_after=40)

para("Tema 20", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=6)
para("DISTRIBUIRANI OKVIR ZA MAŠINSKO UČENJE",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, space_after=4)
para("ZA OBUKU MODELA PREKO ČVOROVA",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, space_after=100)

para("Student:  Vuk Trifunović", align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
     space_after=4)
para("Broj indeksa:  6223", align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
     space_after=60)

para("[Mesto], [datum]", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

pagebreak()

# -------------------------------------------------------------- 2. SADRZAJ --
# Naslov sadrzaja se formatira rucno, a ne kao Heading 1, da se ne bi
# pojavio kao stavka u sopstvenom sadrzaju.
para("Sadržaj", align=WD_ALIGN_PARAGRAPH.LEFT, size=15, bold=True,
     space_after=10)
tocp = doc.add_paragraph()
field(tocp, r'TOC \o "1-2" \h \z \u')
para("(Sadržaj se osvežava u Word-u: desni klik na tabelu iznad → "
     "Update Field → Update entire table.)", size=9, italic=True)

pagebreak()

# ============================================================ 3. TEORIJA ====
h1("3. Uvod i teoretska postavka")

h2("3.1 Uvod")
p("Obuka modela mašinskog učenja svodi se na ponavljanje jednog koraka: "
  "izračunaj gradijent funkcije gubitka nad delom podataka, pomeri parametre u "
  "suprotnom smeru, ponovi. Dok sve staje na jednu mašinu, to je problem "
  "numeričke optimizacije. Čim se posao razdeli na više čvorova, postaje problem "
  "distribuiranog sistema: čvorovi moraju da se usaglase oko zajedničkog stanja, "
  "neko mora da primeti kada čvor prestane da odgovara, a najsporiji učesnik "
  "počinje da određuje tempo svih ostalih.")
p("Ovaj rad opisuje distribuirani okvir za obuku modela preko više čvorova. "
  "Okvir je napisan u Pythonu, direktno nad TCP soketima, bez biblioteka za "
  "razmenu poruka i bez gotovih okvira za mašinsko učenje. Model koji se obučava "
  "je softmax regresija nad MNIST skupom rukom pisanih cifara. Izbor modela je "
  "namerni: predmet rada je raspodela posla, otkrivanje otkaza i posredovanje u "
  "komunikaciji, a ne dubina neuronske mreže. Sve što bi složeniji model doneo "
  "jeste duže računanje po rundi, što bi zamaglilo upravo one mrežne efekte koji "
  "se ovde mere.")
p("Poglavlje 3 postavlja teorijski okvir: paralelizam podataka, arhitekturu "
  "parameter servera, sinhronu i asinhronu obuku, detekciju otkaza, middleware i "
  "interceptore, problem sporog čvora i trajnost stanja. Poglavlje 4 opisuje "
  "konkretnu implementaciju i, što je važnije, meri je. Svaka tvrdnja u poglavlju "
  "4 potkrepljena je brojevima iz baze od trinaest eksperimentalnih pokretanja i "
  "15.684 zabeležene runde obuke.")

h2("3.2 Zašto se obuka modela distribuira")
p("Postoje tri razloga, i vredi ih razdvojiti jer vode ka različitim "
  "arhitekturama.")
bullets([
    "Vreme. Broj operacija po epohi raste sa proizvodom broja uzoraka i broja "
    "parametara. Kada obuka na jednoj mašini traje danima, jedini način da se "
    "skrati jeste da više čvorova radi istovremeno.",
    "Memorija. Model ili skup podataka mogu jednostavno da ne stanu u radnu "
    "memoriju jednog čvora. Tada distribucija nije optimizacija nego uslov "
    "izvodljivosti.",
    "Lokacija podataka. Podaci često već jesu raspoređeni po čvorovima i ne smeju "
    "se centralizovati, bilo zbog količine bilo zbog pravnih ograničenja. Tada se "
    "računanje seli podacima, a ne obrnuto.",
])
p("Rad se bavi prvim slučajem. MNIST staje u memoriju svakog čvora, pa je "
  "distribucija ovde sredstvo za skraćivanje vremena, a ne nužnost. To je "
  "pogodno za merenje, jer omogućava da se distribuirani rezultat direktno uporedi "
  "sa sekvencijalnim referentnim rezultatom nad istim podacima.")

h2("3.3 Paralelizam podataka i paralelizam modela")
p("Postoje dva osnovna načina da se obuka razdeli. Kod paralelizma modela, sam "
  "model se seče na delove i svaki čvor drži svoj deo parametara. Kod paralelizma "
  "podataka, svaki čvor drži celu kopiju modela, ali računa nad različitim "
  "podskupom podataka. Paralelizam modela je neophodan kada model ne staje u "
  "memoriju čvora; paralelizam podataka je jednostavniji i primenjuje se u "
  "ogromnoj većini slučajeva. Ovaj rad koristi paralelizam podataka.")
p("Paralelizam podataka počiva na jednom svojstvu gradijenta. Ako je funkcija "
  "gubitka srednja vrednost po uzorcima, onda je gradijent nad unijom disjunktnih "
  "podskupova težinski prosek gradijenata nad pojedinačnim podskupovima, gde su "
  "težine veličine podskupova:")
formula("ḡ  =  ( n₁·g₁ + n₂·g₂ + … + n_k·g_k ) / ( n₁ + n₂ + … + n_k )")
p("Ova jednakost je razlog zbog kojeg paralelizam podataka uopšte daje tačan "
  "rezultat, a ne samo približan. Iz nje sledi tvrdnja koju poglavlje 4 "
  "eksperimentalno proverava: N sinhronih čvorova sa batch-om B matematički je "
  "ekvivalentno jednom čvoru sa batch-om N·B. Takođe sledi jedno praktično "
  "upozorenje. Ako čvorovi vraćaju srednje vrednosti nad svojim batch-evima, a "
  "batch-evi nisu jednake veličine, obično usrednjavanje bez težina daje pogrešan "
  "rezultat, jer precenjuje doprinos manjih batch-eva. Ta zamka postaje stvarna "
  "čim se uvede balansiranje opterećenja, o čemu govori odeljak 3.8.")

h2("3.4 Arhitektura parameter servera")
p("Parameter server je topologija u kojoj jedan čvor drži jedinu merodavnu kopiju "
  "parametara modela, a radni čvorovi od njega preuzimaju parametre, računaju "
  "gradijente nad svojim podacima i vraćaju ih. Obrazac je opisan i imenovan u "
  "radu Li i saradnika [3], a bio je u upotrebi i ranije, u sistemu DistBelief "
  "[4]. Jedna runda obuke ima četiri koraka: server šalje parametre, radnici "
  "računaju gradijente, radnici vraćaju gradijente, server ih kombinuje i ažurira "
  "parametre.")
p("Glavna alternativa je AllReduce, gde nema centralnog čvora nego radnici "
  "razmenjuju gradijente međusobno po prstenastoj ili stablastoj topologiji. "
  "AllReduce bolje raspoređuje mrežno opterećenje i nema centralnu tačku otkaza, "
  "ali zahteva da svi čvorovi budu prisutni i sinhronizovani, pa lošije podnosi "
  "gubitak čvora i heterogene brzine.")
p("Parameter server ima tri poznata svojstva koja treba imati u vidu. Prvo, "
  "konzistentnost je jednostavna, jer postoji tačno jedno mesto na kojem stanje "
  "živi. Drugo, server je usko grlo propusnosti, jer sav saobraćaj prolazi kroz "
  "njega. Treće, server je centralna tačka otkaza: ako padne, obuka je izgubljena "
  "osim ako se stanje periodično ne snima na trajni medijum. Odeljak 3.9 i "
  "odeljak 4.12 bave se upravo tim trećim svojstvom.")

h2("3.5 Sinhrona i asinhrona obuka")
p("Kada server dobija gradijente od više radnika, mora da odluči kada da ažurira "
  "parametre. Postoje dva odgovora.")
p("U sinhronom režimu server čeka da svi aktivni radnici pošalju gradijent za "
  "tekuću rundu, pa tek onda kombinuje i ažurira. To je barijera, u literaturi "
  "poznata kao bulk synchronous parallel model. Svi radnici u svakoj rundi računaju "
  "iz iste verzije parametara, pa je rezultat determinisan i identičan "
  "sekvencijalnom računanju sa odgovarajuće većim batch-om. Cena je što svaka "
  "runda traje onoliko koliko traje najsporiji radnik.")
p("U asinhronom režimu server ažurira parametre odmah po prijemu svakog "
  "pojedinačnog gradijenta i istog trenutka vraća radniku nove parametre. Nema "
  "barijere, pa niko nikoga ne čeka. Cena je zastarelost gradijenata (engl. "
  "staleness): dok je radnik računao, server je već primenio gradijente drugih "
  "radnika, pa gradijent koji stiže više ne odgovara trenutnom stanju parametara. "
  "Zastarelost se meri kao razlika između broja tekuće runde na serveru i broja "
  "runde iz koje su parametri preuzeti.")
p("Asinhroni pristup popularizovali su Hogwild! [5] i Downpour SGD [4]. Oba rada "
  "pokazuju da umerena zastarelost u praksi ne sprečava konvergenciju, iako "
  "narušava teorijske garancije. Reč je o klasičnom kompromisu između "
  "konzistentnosti i propusnosti, i odeljak 4.8 ga kvantifikuje na izmerenim "
  "podacima.")

h2("3.6 Detekcija otkaza")
p("U distribuiranom sistemu čvor može da prestane da odgovara na više načina: "
  "proces se ugasi, mašina se zamrzne, mreža se podeli. Sa stanovišta posmatrača "
  "sa druge strane veze, ti slučajevi nisu isti, i to je važnije nego što deluje.")
p("Teorijski rezultat koji stoji iza svega jeste da se u potpuno asinhronoj mreži, "
  "u kojoj poruke mogu kasniti proizvoljno dugo, spor čvor ne može razlikovati od "
  "mrtvog. Iz toga sledi nemogućnost konsenzusa u prisustvu i jednog otkaza (FLP "
  "rezultat), a Chandra i Toueg [6] pokazuju da se problem zaobilazi uvođenjem "
  "nepouzdanih detektora otkaza, koji smeju da greše. Detektor se opisuje sa dva "
  "svojstva: potpunost (svaki mrtav čvor pre ili kasnije bude osumnjičen) i "
  "tačnost (živ čvor ne bude osumnjičen). U praksi se potpunost postiže lako, a "
  "tačnost se žrtvuje.")
p("Praktični detektor zasniva se na isteku vremena. Odabir tog vremena je "
  "suštinski kompromis, a ne detalj podešavanja: kratak tajmaut izbacuje spore ali "
  "žive čvorove i time gubi koristan rad, dugačak tajmaut odlaže oporavak i drži "
  "ceo klaster zaustavljenim. Ne postoji vrednost koja je ispravna u opštem "
  "slučaju.")
p("Postoji, međutim, i drugi signal, koji je tačan i trenutan kada je dostupan. "
  "Ako operativni sistem na drugoj strani zatvori TCP vezu, čitanje sa soketa "
  "vraća prazan rezultat i time se pouzdano zna da proces više ne postoji. Taj "
  "signal ne zahteva nikakvo pogađanje. Ograničenje je što ga ima samo za otkaze "
  "procesa, ne i za zamrznute čvorove ili prekinutu mrežu, gde veza formalno "
  "ostaje otvorena. Zato ozbiljna implementacija koristi oba signala, a poglavlje "
  "4.9 meri koliko se razlikuju po ceni.")

h2("3.7 Middleware i interceptori")
p("Middleware je sloj softvera koji stoji između aplikacije i mreže i preuzima "
  "poslove koji se ponavljaju u svakoj distribuiranoj aplikaciji: serijalizaciju, "
  "kompresiju, merenje, autentifikaciju. Interceptor je, prema Tanenbaumu i Van "
  "Steenu [1], objekat koji se nalazi na putu poruke i sme da je pregleda ili "
  "izmeni, tako da ni pošiljalac ni primalac ne znaju da postoji. Isti obrazac "
  "javlja se kao servlet filter u Javi, kao interceptor u gRPC-u i kao middleware "
  "u većini veb okvira.")
p("Praktična korist je razdvajanje odgovornosti. Kompresija, brojanje bajtova i "
  "veštačko kašnjenje mogu se uključivati i isključivati bez ijedne izmene u "
  "petlji za obuku. Interceptori se ulančavaju po obrascu lanca odgovornosti, i "
  "tu postoji jedno pravilo koje se lako previdi: lanac mora biti simetričan. Ako "
  "se pri slanju primenjuju redom od prvog ka poslednjem, pri prijemu se moraju "
  "primeniti obrnutim redosledom, jer je reč o kompoziciji funkcija koja se mora "
  "odmotati suprotno od načina na koji je namotana.")

h2("3.8 Problem sporog čvora i balansiranje opterećenja")
p("Spor čvor (engl. straggler) je čvor koji zaostaje za ostalima. U sinhronoj "
  "obuci njegov uticaj je nesrazmeran: pošto barijera čeka sve, jedan čvor koji "
  "je tri puta sporiji čini ceo klaster tri puta sporijim, dok preostali čvorovi "
  "stoje besposleno. To besposleno vreme je merljivo i predstavlja direktan "
  "gubitak.")
p("Uzroci su raznovrsni: slabiji hardver, deljenje mašine sa drugim poslovima, "
  "termalno usporavanje, neravnomerna raspodela podataka. Rešenja se dele u dve "
  "grupe. Prva je redundantno izvršavanje, poznato iz MapReduce-a [7], gde se "
  "isti zadatak dodeljuje i rezervnom čvoru pa se uzima brži rezultat. Druga je "
  "prilagođavanje količine posla izmerenoj brzini čvora, tako da svi završe "
  "otprilike istovremeno. Ovaj rad koristi drugi pristup.")
p("Prilagođavanje ima jedno ograničenje koje treba navesti unapred. Balanser može "
  "da pomeri samo onaj deo cene koji raste sa količinom posla. Ako čvor ima "
  "fiksni trošak po rundi, nezavisan od veličine batch-a, smanjivanje batch-a "
  "neće pomoći. Odeljak 4.11 zato razlikuje ta dva izvora sporosti i meri ih "
  "odvojeno.")

h2("3.9 Trajnost stanja i atomski upis")
p("Pošto server drži jedinu kopiju parametara, njegov pad znači gubitak celokupne "
  "obuke. Standardno rešenje jeste periodično snimanje stanja na trajni medijum "
  "(engl. checkpointing), čime se gubitak ograničava na rad obavljen od poslednjeg "
  "snimka.")
p("Snimanje ima suptilnost koja je važnija od same ideje. Snimak koji je upisan "
  "do pola u trenutku pada gori je nego da ga uopšte nema, jer se pri sledećem "
  "pokretanju čita skraćen fajl i greška se manifestuje na način koji podseća na "
  "grešku u modelu, a ne na oštećen fajl. Rešenje je atomski upis: podaci se prvo "
  "upisuju u privremeni fajl u istom direktorijumu, zatim se pozivom fsync "
  "obezbeđuje da su stvarno stigli na disk, i tek onda se privremeni fajl "
  "preimenuje preko odredišnog. Operacija preimenovanja je po POSIX standardu "
  "atomska unutar jednog fajl sistema, pa čitalac uvek vidi ili ceo stari ili ceo "
  "novi snimak, nikada mešavinu. Zahtev da fajlovi budu u istom direktorijumu nije "
  "formalnost: preko granice fajl sistema preimenovanje se svodi na kopiranje i "
  "atomičnost se gubi.")

pagebreak()

# ======================================================== 4. STUDIJA SLUCAJA ==
h1("4. Studija slučaja: implementacija i merenja")

h2("4.1 Cilj i obim")
p("Implementiran je distribuirani okvir za obuku modela po topologiji parameter "
  "servera, nad golim TCP soketima. Cilj nije bio napraviti brz sistem za "
  "mašinsko učenje, nego sistem u kojem se svaki mehanizam iz poglavlja 3 može "
  "uključiti, isključiti i izmeriti.")
p("Model je softmax regresija: 784 ulaza, 10 klasa, uz dodatni bias član, dakle "
  "matrica parametara dimenzije 785 × 10, odnosno 7.850 parametara ili 62,8 KB u "
  "dvostrukoj preciznosti. Skup podataka je MNIST, od kojeg se koristi 10.000 "
  "uzoraka za obuku i 2.000 izdvojenih uzoraka za proveru. Permutacija skupa "
  "vezana je za fiksno seme, pa je podela na skup za obuku i skup za proveru ista "
  "pri svakom pokretanju.")
p("Sve nadogradnje opisane u nastavku podrazumevano su isključene. Pokretanje bez "
  "dodatnih zastavica daje istu putanju izvršavanja kao osnovna verzija sistema, "
  "što je bio uslov da rezultati iz ranijih merenja ostanu uporedivi sa kasnijim.")
p("Kod je podeljen u dva paketa: N1 nosi klijent-server sloj, a N2 middleware i "
  "skladištenje. Tabela 1 daje pregled modula i njihovih uloga.")

table("Tabela 1. Struktura projekta i uloga pojedinačnih modula.",
      ["Modul", "Sloj", "Uloga"],
      [["N1/common.py", "model", "Učitavanje MNIST-a, softmax, gradijent, evaluacija"],
       ["N1/net.py", "komunikacija", "Uokvirivanje poruka preko TCP-a, poziv lanca interceptora"],
       ["N1/server.py", "komunikacija", "Parameter server: petlja, barijera, detektor otkaza, dodela posla"],
       ["N1/worker.py", "komunikacija", "Radni čvor i namerno ubrizgavanje kvara"],
       ["N2/interceptors.py", "middleware", "Lanac interceptora: float32, deflate, metrike, kašnjenje"],
       ["N2/balancer.py", "middleware", "Procena brzine radnika i dodela veličine batch-a"],
       ["N2/checkpoint.py", "middleware", "Atomsko snimanje i učitavanje parametara"],
       ["N2/store.py", "baza", "SQLite šema, upis metrika, čitanje rezultata"],
       ["DOC/baseline.py", "merenje", "Sekvencijalni referentni run, bez mreže"],
       ["DOC/plot.py", "merenje", "Generisanje svih sedam figura iz baze"]],
      [2.0, 1.2, 5.0])

p("Podela na pakete N1 i N2 prati podelu iz tabele za ocenjivanje: N1 je "
  "klijent-server sloj, N2 je middleware sloj. Podela nije samo organizaciona. "
  "Modul net.py ne zna ništa o obuci, interceptors.py ne zna ništa o soketima, a "
  "balancer.py nema nijedan uvoz i ne zna ni za mrežu ni za numpy. Svaki od njih "
  "može se ispitivati nezavisno od ostatka sistema.")

h2("4.2 Komunikacioni sloj")
p("TCP je tok bajtova, a ne tok poruka. Jedan poziv slanja može stići kao "
  "nekoliko poziva čitanja, a dva slanja mogu stići kao jedno čitanje. Granica "
  "poruke ne postoji i mora se uvesti na aplikativnom nivou. Ovde se koristi "
  "prefiks dužine: svaka poruka počinje sa četiri bajta koji nose broj bajtova "
  "koji slede, zapisan u mrežnom redosledu bajtova, pa protokol ne zavisi od "
  "arhitekture mašine.")
p("Druga strana istog problema je čitanje. Poziv recv sme da vrati manje bajtova "
  "nego što je traženo, i to nije greška nego uobičajeno ponašanje. Zato se čita "
  "u petlji, dok se ne prikupi tačno onoliko bajtova koliko je najavljeno:")
code([
    "def recv_exactly(sock, n):",
    "    buf = bytearray()",
    "    while len(buf) < n:",
    "        chunk = sock.recv(n - len(buf))",
    "        if not chunk:",
    "            return None",
    "        buf += chunk",
    "    return bytes(buf)",
])
p("Ako se ovo napiše kao običan poziv recv, kod radi ispravno na malim porukama i "
  "otkazuje na velikim. Gradijent u ovom sistemu je 62,8 KB, što je sasvim "
  "dovoljno da se problem ispolji. Vrednost None znači da je druga strana uredno "
  "zatvorila vezu, a to je istovremeno i prvi signal detektora otkaza iz odeljka "
  "3.6.")
p("Protokol ima svega dva tipa poruke od radnika ka serveru i jedan odgovor u "
  "suprotnom smeru; tabela 2 daje njihov format.")

table("Tabela 2. Format poruka aplikativnog protokola.",
      ["Poruka", "Smer", "Polja", "Značenje"],
      [["PULL", "radnik → server", "type, worker",
        "Registracija radnika u klaster"],
       ["PUSH", "radnik → server", "type, worker, grads, loss, n, round",
        "Gradijent nad dodeljenim batch-om i verzija parametara iz koje je računat"],
       ["odgovor", "server → radnik", "weights, round, assign | stop",
        "Nova verzija parametara i opseg redova za sledeću rundu, ili nalog za prekid"]],
      [1.1, 1.5, 2.6, 3.2])

p("Polje round u poruci PUSH nosi verziju parametara iz koje je gradijent "
  "izračunat. Server iz razlike svoje tekuće runde i te vrednosti dobija "
  "zastarelost gradijenta, o kojoj govori odeljak 4.8. Polje n nosi stvarnu "
  "veličinu batch-a i neophodno je za težinsko usrednjavanje iz odeljka 4.11.")

h2("4.3 Parameter server: petlja i barijera")
p("Server je jednonitni i koristi multipleksiranje nad svim otvorenim vezama:")
code([
    "readable, _, _ = select.select([lsock] + list(conns), [], [], 0.5)",
])
p("Izbor jedne niti nije posledica jednostavnosti nego namere. Parametri modela "
  "postoje u tačno jednom primerku i menjaju se u svakoj rundi. Sa više niti "
  "svaka bi dirala isto stanje, pa bi bilo neophodno zaključavanje; ovako postoji "
  "jedan tok izvršavanja i trka za podacima nije moguća. Tajmaut od pola sekunde "
  "postoji da bi se detektor otkaza izvršavao i onda kada nema nikakvog "
  "saobraćaja, jer zamrznut čvor po definiciji ništa ne šalje.")
p("Barijera u sinhronom režimu razrešava se kada skup radnika koji su odgovorili "
  "postane jednak skupu aktivnih radnika. U toj tački server kombinuje gradijente "
  "težinskim prosekom, ažurira parametre i otpušta radnike u novu rundu. Tri "
  "detalja u toj logici nisu očigledna i vredi ih izdvojiti.")
p("Prvo, detektor otkaza meri vreme od trenutka otpuštanja radnika, a ne od "
  "poslednje primljene poruke. Razlog je što radnik koji čeka na barijeri po "
  "definiciji ćuti. Kada bi se merilo vreme od poslednje poruke, izbacivanje "
  "jednog radnika povuklo bi za sobom izbacivanje svih preostalih. Tajmaut ovde "
  "znači vreme otkad je radniku dodeljen posao, a ne vreme otkad se javio.")
p("Drugo, detektor se ne izvršava dok se klaster ne formira, jer bi inače prvi "
  "radnik koji se prijavi bio izbačen dok se čeka poslednji. Iz istog razloga se i "
  "merenje vremena pokreće tek kada su svi radnici prisutni, da rezultati ne bi "
  "sadržali vreme pokretanja procesa.")
p("Treće, poruka koja stigne od već izbačenog radnika mora da prekine vezu. Radnik "
  "može biti izbačen po tajmautu a da je i dalje živ, i ako bi mu se dozvolilo da "
  "uđe u skup onih koji su odgovorili, barijera bi se zaključala trajno, jer taj "
  "skup više nikada ne bi bio jednak skupu aktivnih radnika.")

h2("4.4 Eksperimentalna postavka")
p("Svi eksperimenti pokrenuti su na jednoj mašini, pri čemu su server i svaki "
  "radnik zasebni procesi koji komuniciraju preko lokalne mrežne petlje. "
  "Zajednički parametri su: korak učenja 0,05, batch od 32 uzorka po radniku, i "
  "budžet od 100.000 obrađenih uzoraka, što odgovara deset prolaza kroz skup za "
  "obuku. Server je vlasnik uslova zaustavljanja, pa svi runovi troše isti broj "
  "uzoraka bez obzira na broj radnika, što ih čini uporedivim.")
p("Rezultati svih pokretanja upisuju se u SQLite bazu opisanu u odeljku 4.13. "
  "Brojevi navedeni u nastavku očitani su iz te baze, a ne iz ispisa na konzoli.")

h2("4.5 Ispravnost sinhronizacije")
p("Prvo pitanje na koje sistem mora da odgovori nije koliko je brz, nego da li "
  "uopšte računa ono što tvrdi da računa. Prema jednakosti iz odeljka 3.3, četiri "
  "sinhrona radnika sa batch-om od 32 uzorka moraju dati isti rezultat kao jedan "
  "čvor sa batch-om od 128 uzoraka. Zato je implementiran i sekvencijalni "
  "referentni run, potpuno bez mreže, koji služi kao merilo.")

figure("fig1_correctness.png",
       "Slika 1. Ispravnost sinhronizacije. Sekvencijalni run sa batch-om od 128 "
       "uzoraka i sinhroni run sa četiri radnika po 32 uzorka daju istu krivu "
       "tačnosti u funkciji broja obrađenih uzoraka.", 5.1)

p("Slika 1 pokazuje da se krive poklapaju. Poklapanje je, međutim, jače nego što "
  "se sa grafika vidi. Poređenjem vrednosti funkcije gubitka na istim tačkama, "
  "prvih 78 rundi, odnosno ceo prvi prolaz kroz skup za obuku, daje brojčano "
  "identične vrednosti do pete decimale. Distribuirano računanje ovde nije samo "
  "približno tačno, nego reprodukuje sekvencijalno računanje.")
p("Posle prvog prolaza javlja se malo razilaženje. Uzrok nije greška u "
  "sinhronizaciji nego sastav batch-eva na granici epohe: skup od 10.000 uzoraka "
  "nije deljiv sa 128, pa poslednja runda svake epohe ima skraćen batch, a taj "
  "skraćeni batch u sekvencijalnom slučaju obrađuje jedan čvor, dok se u "
  "distribuiranom slučaju deli na četiri. Od te tačke krive se blago razdvajaju, "
  "ali ostaju u granicama od 0,0055 apsolutne razlike u tačnosti, uz finalne "
  "vrednosti 0,8875 i 0,8870. Zaključak je da je sinhronizacija ispravna, i to je "
  "preduslov da bilo koje merenje brzine u nastavku uopšte ima smisla.")

h2("4.6 Skalabilnost")
p("Sa potvrđenom ispravnošću, sledeće pitanje je šta se dobija povećanjem broja "
  "čvorova. Odgovor zavisi od toga šta se meri, i upravo to je najzanimljiviji "
  "nalaz ovog odeljka.")

figure("fig2_speedup.png",
       "Slika 2. Skalabilnost. Levo: konvergencija u realnom vremenu za jedan, dva "
       "i četiri radnika. Desno: dve vrste ubrzanja u odnosu na jednog radnika.")

table("Tabela 3. Rezultati skaliranja pri istom budžetu od 100.000 uzoraka.",
      ["Radnika", "Rundi", "Vreme (s)", "Uzoraka/s", "Ubrzanje", "Vreme do 0,869 (s)", "Finalna tačnost"],
      [["1", "3.130", "19,21", "5.205", "1,00×", "1,98", "0,9070"],
       ["2", "1.565", "10,36", "9.651", "1,85×", "2,09", "0,8990"],
       ["4", "783", "5,46", "18.336", "3,52×", "2,02", "0,8870"]],
      [1.0, 1.0, 1.1, 1.1, 1.0, 1.5, 1.4],
      align_right=(0, 1, 2, 3, 4, 5, 6))

p("Propusnost, izražena brojem obrađenih uzoraka u sekundi, raste skoro linearno: "
  "1,85 puta na dva radnika i 3,52 puta na četiri. Odstupanje od idealnog "
  "linearnog rasta potiče od serijalizacije na serveru, koji sve gradijente "
  "obrađuje u jednoj niti.")
p("Vreme potrebno da se dostigne zadata tačnost, međutim, ne poboljšava se uopšte: "
  "1,98 sekundi na jednom radniku prema 2,02 sekunde na četiri. Razlog je što N "
  "radnika pri nepromenjenom koraku učenja daje N puta veći globalni batch, pa "
  "svaki korak optimizacije nosi srazmerno manje napretka. Klaster obrađuje više "
  "podataka u sekundi, ali mu treba više podataka za isti pomak. To se vidi i u "
  "poslednjoj koloni tabele 3: pri istom budžetu uzoraka, više radnika znači manje "
  "koraka optimizacije, pa i nižu finalnu tačnost, 0,9070 prema 0,8870.")
p("Ovaj nalaz nije nedostatak implementacije nego poznato svojstvo paralelizma "
  "podataka. Praktičan zaključak je da se sa brojem čvorova mora skalirati i korak "
  "učenja, inače se dobija sistem koji brže troši podatke bez ubrzanja obuke. "
  "Prikazivanje samo krive propusnosti prećutalo bi ovu cenu, a prikazivanje samo "
  "krive vremena do tačnosti ostavilo bi utisak da paralelizacija ne radi. Zbog "
  "toga desni deo slike 2 prikazuje obe.")

h2("4.7 Komunikacioni trošak")
p("Podatak koji objašnjava zašto skaliranje nije idealno vidi se iz količine "
  "prenetih podataka. Za obuku modela čiji parametri zauzimaju 62,8 KB, kroz mrežu "
  "je prošlo 197,49 MB.")

figure("fig5_bandwidth.png",
       "Slika 3. Komunikacioni trošak: količina podataka primljena na serveru u "
       "funkciji broja obrađenih uzoraka, za jednog, dva i četiri radnika. Tri "
       "krive se potpuno poklapaju.", 5.1)

p("Na slici 3 tri krive se poklapaju, i to nije greška u crtanju nego rezultat. "
  "Veličina poruke sa gradijentom ne zavisi od veličine batch-a, jer gradijent "
  "uvek ima dimenziju parametara modela. Iz toga sledi da komunikacioni trošak po "
  "obrađenom uzorku ne zavisi od broja čvorova, nego isključivo od veličine "
  "batch-a po čvoru. Izmereno, radnik pošalje 63.057 bajtova po rundi, što pri "
  "batch-u od 32 uzorka daje 1.962 bajta po uzorku, isto za sve tri "
  "konfiguracije.")
p("Praktična posledica je da se komunikacija ne smanjuje smanjivanjem broja "
  "čvorova, nego povećanjem batch-a po čvoru ili kompresijom poruka. Drugi put "
  "obrađuje odeljak 4.10.")

h2("4.8 Sinhroni i asinhroni režim")
p("Poređenje ova dva režima zahteva pažnju pri postavljanju pitanja. Pri istom "
  "budžetu uzoraka, sinhroni run sa četiri radnika izvršava 783 runde, jer se u "
  "svakoj rundi troši 128 uzoraka odjednom. Asinhroni run sa četiri radnika "
  "izvršava 3.130 rundi, jer se parametri ažuriraju posle svakog pojedinačnog "
  "gradijenta nad 32 uzorka. To su različite količine posla optimizacije i ne "
  "porede se neposredno.")

figure("fig3_sync_vs_async.png",
       "Slika 4. Sinhroni i asinhroni režim. Levo: tačnost u funkciji stvarnog "
       "vremena. Desno: raspodela zastarelosti gradijenata u asinhronom režimu.")

p("Ispravno poređenje jeste sa sinhronim runom koji izvršava isti broj koraka "
  "optimizacije, a to je run sa jednim radnikom. Oba izvršavaju 3.130 rundi i "
  "dostižu praktično istu tačnost, 0,9070 sinhrono i 0,9075 asinhrono. Razlika je "
  "u vremenu: 19,21 sekundi prema 7,34 sekunde, dakle 2,62 puta brže. To je "
  "stvarni doprinos asinhronog režima. Četiri čvora obavljaju isti posao "
  "optimizacije kao jedan, ali paralelno.")
p("Cena je zastarelost gradijenata. U sinhronom režimu ona je nula u svakoj rundi, "
  "po definiciji barijere. U asinhronom režimu prosečna zastarelost iznosi 2,97 "
  "rundi, sa maksimumom od 7. Desni deo slike 4 pokazuje raspodelu: najčešća "
  "vrednost je 3, što odgovara broju ostalih radnika u klasteru, a to je i "
  "očekivano, jer dok jedan radnik računa, preostala trojica u proseku stignu da "
  "pošalju po jedan gradijent.")

table("Tabela 4. Poređenje sinhronog i asinhronog režima.",
      ["Svojstvo", "Sinhrono", "Asinhrono"],
      [["Ažuriranje parametara", "tek kada svi pošalju gradijent", "odmah po prijemu svakog gradijenta"],
       ["Barijera", "postoji", "ne postoji"],
       ["Zastarelost gradijenata", "0 u svakoj rundi", "prosečno 2,97, najviše 7"],
       ["Tempo određuje", "najsporiji radnik", "svaki radnik nezavisno"],
       ["Rundi pri 100.000 uzoraka", "783 (4 radnika)", "3.130 (4 radnika)"],
       ["Vreme pri 3.130 rundi", "19,21 s (1 radnik)", "7,34 s (4 radnika)"],
       ["Tačnost pri 3.130 rundi", "0,9070", "0,9075"],
       ["Determinisanost", "potpuna", "zavisi od redosleda dolaska poruka"]],
      [2.4, 2.3, 2.8])

p("Tabela 4 sažima kompromis. Sinhroni režim daje determinisan rezultat koji se "
  "može uporediti sa sekvencijalnim, po cenu čekanja na najsporiji čvor. Asinhroni "
  "režim uklanja čekanje, ali gubi determinisanost, jer rezultat zavisi od "
  "redosleda kojim poruke stignu. Za ovaj model i ovaj nivo zastarelosti gubitak "
  "tačnosti nije izmeren, što se slaže sa nalazima iz literature [4][5]. Kod "
  "većih klastera zastarelost raste sa brojem čvorova i tada prestaje da bude "
  "bezopasna.")

h2("4.9 Tolerancija otkaza")
p("Za proveru detekcije otkaza implementirana su dva različita kvara, koja "
  "odgovaraju dvama signalima iz odeljka 3.6. Prvi je pad procesa: radnik u "
  "zadatoj rundi poziva neposredan izlaz, bez čišćenja, pa operativni sistem "
  "zatvara soket. Drugi je zamrzavanje: radnik ulazi u beskonačno spavanje, veza "
  "ostaje otvorena, a poruke prestaju. Reč je o dva različita eksperimenta i ne "
  "smeju se predstaviti kao jedan.")

figure("fig4_fault_tolerance.png",
       "Slika 5. Tolerancija otkaza. Levo: obuka se nastavlja posle gubitka čvora "
       "u rundi 150. Desno: trajanje runde u logaritamskoj razmeri, gde se vidi "
       "cena detekcije.")

table("Tabela 5. Poređenje dva mehanizma otkaza i cene njihove detekcije.",
      ["", "Pad procesa", "Zamrznut čvor"],
      [["Mehanizam kvara", "neposredan izlaz iz procesa", "beskonačno spavanje"],
       ["Stanje TCP veze", "zatvorena", "otvorena"],
       ["Signal detekcije", "čitanje vraća prazan rezultat", "istek tajmauta"],
       ["Podešen tajmaut", "8,0 s", "5,0 s"],
       ["Razlog upisan u bazu", "connection closed", "timeout"],
       ["Trajanje runde 150", "0,00600 s", "0,00542 s"],
       ["Trajanje runde 151", "0,01349 s", "5,09854 s"],
       ["Obuka nastavljena do runde", "794", "794"],
       ["Finalna tačnost", "0,8895", "0,8895"]],
      [2.3, 2.3, 2.3])

p("Levi deo slike 5 pokazuje da obuka preživljava gubitak čvora: kriva tačnosti "
  "prolazi kroz rundu 150 bez vidljivog prekida i oba runa dolaze do kraja "
  "budžeta uzoraka sa istom finalnom tačnošću od 0,8895. Barijera se prilagođava "
  "smanjenom skupu aktivnih radnika, a preostala tri čvora nastavljaju rad.")
p("Desni deo slike 5 pokazuje ono što je zapravo predmet merenja, a to je cena "
  "detekcije. Kod pada procesa runda u kojoj je otkaz otkriven traje 13,5 "
  "milisekundi, dakle jedva primetno duže od uobičajenih 6 milisekundi. Kod "
  "zamrznutog čvora ista runda traje 5,0985 sekundi, što je podešeni tajmaut "
  "uvećan za trajanje runde. Odnos je približno 378 puta.")
p("Ta razlika je suština problema iz odeljka 3.6. Signal zatvorene veze je tačan i "
  "besplatan, ali postoji samo zato što otkaz ovde pogađa proces na istoj mašini. "
  "Pravi distribuirani sistem, kod kojeg otkazuje udaljeni čvor ili mrežna veza, "
  "za takve slučajeve raspolaže isključivo drugim signalom, a njegova cena je "
  "podešeni tajmaut. Skraćivanje tajmauta ne rešava problem, nego ga zamenjuje "
  "drugim: izbacivanjem sporih ali živih čvorova i gubitkom njihovog rada.")

h2("4.10 Middleware: lanac interceptora")
p("Middleware sloj implementiran je kao lanac interceptora sa dva nivoa kuka. "
  "Objektni nivo vidi poruku kao rečnik, pre serijalizacije, i koristi ga "
  "interceptor koji gradijente iz dvostruke prevodi u jednostruku preciznost. "
  "Bajtni nivo vidi serijalizovani sadržaj i ne zna šta bajtovi znače; koristi ga "
  "deflate kompresija i brojač saobraćaja. Dva nivoa postoje zato što dve korisne "
  "transformacije prirodno pripadaju različitim slojevima.")
p("Lanac se primenjuje simetrično, kako nalaže odeljak 3.7: pri slanju od prvog ka "
  "poslednjem, pri prijemu obrnuto. Ako se redosled pri prijemu ne obrne, "
  "dekompresiji se prosleđuje serijalizovani sadržaj koji ona ne ume da pročita.")
p("Postoji i razlika koju treba istaći. Bajtni interceptori moraju biti uključeni "
  "na obe strane veze, jer se ono što je jedna strana komprimovala mora "
  "dekomprimovati na drugoj. Objektni interceptori mogu biti jednostrani: radnik "
  "koji šalje gradijent u jednostrukoj preciznosti ne zahteva ništa od servera, "
  "jer numpy pri prvoj aritmetičkoj operaciji sam prevodi uži tip u širi. Zbog "
  "toga zastavica za jednostruku preciznost postoji samo na radniku, a zastavica "
  "za kompresiju mora se navesti na oba procesa.")

figure("fig7_interceptor_bandwidth.png",
       "Slika 6. Efekat interceptora za kompresiju na količinu podataka primljenu "
       "na serveru.", 5.1)

table("Tabela 6. Efekat interceptora pri istom budžetu uzoraka (783 runde).",
      ["Konfiguracija", "Lanac na radniku", "Primljeno (MB)", "Odnos", "Vreme (s)"],
      [["Bez kompresije", "Chain([metrics])", "197,49", "1,000", "5,46"],
       ["float32 gradijenti", "Chain([float32, metrics])", "99,26", "0,503", "5,31"],
       ["float32 + deflate 6", "Chain([float32, zlib, metrics])", "54,67", "0,277", "12,92"]],
      [1.8, 2.6, 1.4, 0.9, 1.0],
      align_right=(2, 3, 4))

p("Slika 6 i tabela 6 prikazuju isto merenje, jednom kao krive u vremenu i jednom "
  "kao krajnje vrednosti. Prelazak na jednostruku preciznost prepolovljava "
  "saobraćaj, sa 197,49 MB na 99,26 MB, i to bez ikakvog uticaja na tačnost, koja "
  "u oba slučaja iznosi 0,8870. Vreme izvršavanja se čak neznatno smanjuje. To je "
  "čist dobitak.")
p("Deflate kompresija dodatno smanjuje saobraćaj na 54,67 MB, ali vreme "
  "izvršavanja raste sa 5,46 na 12,92 sekunde, dakle 2,37 puta. Merenja "
  "prikupljena samim interceptorom objašnjavaju zašto: na strani servera 197,48 MB "
  "sirovih podataka svedeno je na 160,07 MB, što je odnos od svega 0,81, uz "
  "utrošenih 6,49 sekundi procesorskog vremena. Gust niz brojeva u pokretnom "
  "zarezu blizu je nekompresibilnog, a na lokalnoj mrežnoj petlji prenos je "
  "gotovo besplatan, pa uštedu u bajtovima nadmašuje cena računanja.")
p("Ovaj rezultat je jedini razlog zbog kojeg interceptor uopšte prikuplja "
  "statistiku o odnosu kompresije i utrošenom procesorskom vremenu. Zaključak nije "
  "bio poznat unapred, nego je izmeren, i suprotan je očekivanju da kompresija "
  "uvek pomaže. Na stvarnoj mreži sa ograničenim propusnim opsegom odnos bi bio "
  "drugačiji, ali upravo to je poenta: odluka o kompresiji zavisi od odnosa cene "
  "procesora i cene prenosa, i mora se meriti u okruženju u kojem sistem radi.")

h2("4.11 Balansiranje opterećenja")
p("Za proveru ponašanja u prisustvu sporog čvora pokrenuta su dva runa sa istim "
  "kvarom i različitim odgovorom sistema. U oba je prvi radnik podešen da troši "
  "četiri puta više vremena po uzorku od ostalih. Prvi run koristi statičku "
  "podelu, gde svi dobijaju jednake batch-eve; drugi koristi dinamičko "
  "balansiranje.")
p("Balanser održava eksponencijalno otežanu procenu vremena po uzorku za svakog "
  "radnika. Merenje se obavlja na jedinom mestu gde je moguće, između trenutka "
  "otpuštanja radnika sa barijere i trenutka dolaska njegovog gradijenta. Globalni "
  "batch se zatim deli srazmerno izmerenoj brzini.")

figure("fig6_load_balancing.png",
       "Slika 7. Balansiranje opterećenja. Levo: veličina batch-a koju server "
       "dodeljuje svakom radniku kroz runde. Desno: prosečno čekanje na barijeri "
       "pri statičkoj podeli i pri dinamičkom balansiranju.")

table("Tabela 7. Raspodela posla i vreme odziva po radniku, prosek za runde nakon 20.",
      ["Radnik", "Statička podela: batch", "Statička podela: vreme (s)",
       "Balansiranje: batch", "Balansiranje: vreme (s)"],
      [["worker-1 (spor)", "31,9", "0,01417", "10,0", "0,00519"],
       ["worker-2", "31,9", "0,00467", "39,3", "0,00514"],
       ["worker-3", "32,0", "0,00462", "40,0", "0,00516"],
       ["worker-4", "32,0", "0,00460", "38,3", "0,00517"]],
      [1.7, 1.4, 1.6, 1.4, 1.6],
      align_right=(1, 2, 3, 4))

p("Tabela 7 pokazuje i problem i rešenje. Pri statičkoj podeli svi radnici dobijaju "
  "po 32 uzorka, ali spor radnik troši 14,2 milisekunde dok ostali troše oko 4,6, "
  "pa preostala tri čvora provode oko dve trećine svakog ciklusa čekajući. Pri "
  "dinamičkom balansiranju spor radnik dobija 10 uzoraka, ostali oko 39, i sva "
  "četiri vremena odziva izjednačuju se na približno 5,2 milisekunde.")
p("Izmerena raspodela poklapa se sa teorijskom. Pri odnosu brzina 1 : 4 : 4 : 4 i "
  "globalnom batch-u od 128 uzoraka, srazmerna podela daje 128/13 ≈ 9,85 za sporog "
  "radnika i 4·128/13 ≈ 39,4 za ostale. Izmereno je 10,0 odnosno 38,3 do 40,0.")
p("Efekat na performanse je znatan. Prosečno čekanje na barijeri smanjeno je sa "
  "0,00993 na 0,00066 sekundi, dakle petnaest puta. Prosečno trajanje runde "
  "smanjeno je sa 0,01568 na 0,00676 sekundi, a ukupno vreme izvršavanja sa 12,27 "
  "na 5,32 sekunde, odnosno 2,31 puta.")
p("Dva svojstva implementacije zaslužuju objašnjenje. Prvo, globalni batch ostaje "
  "konstantan: klaster u svakoj rundi troši tačno onoliko uzoraka koliko bi "
  "potrošio i bez balansiranja, a greška zaokruživanja se pripisuje najbržem "
  "radniku da bi zbir bio tačan. Zahvaljujući tome statički i dinamički run "
  "uporedivi su pri istom broju obrađenih uzoraka.")
p("Drugo, i važnije, prosek gradijenata mora postati težinski. Svaki radnik vraća "
  "srednju vrednost nad svojim batch-em, a usrednjavanje srednjih vrednosti sa "
  "jednakim težinama ispravno je samo ako su batch-evi jednaki. Sa nejednakim "
  "batch-evima jedino težinski prosek iz odeljka 3.3 daje srednju vrednost nad "
  "unijom batch-eva:")
code([
    "ns = np.array([m[2] for m in buf_meta], dtype=np.float64)",
    "combined = np.tensordot(ns, stacked, axes=(0, 0)) / ns.sum()",
])
p("Obično usrednjavanje bez težina precenilo bi doprinos manjih batch-eva sporih "
  "radnika i tiho promenilo funkciju koja se optimizuje. Zahvaljujući težinskom "
  "proseku, ekvivalencija sa sekvencijalnim računanjem iz odeljka 4.5 važi i uz "
  "uključeno balansiranje.")
p("Napokon, treba navesti granicu ove tehnike. Balanser može da pomeri samo onaj "
  "deo cene koji raste sa količinom posla. Zato su u sistemu predviđena dva "
  "odvojena načina simulacije sporosti: fiksno kašnjenje po rundi, nezavisno od "
  "veličine batch-a, i kašnjenje po uzorku. Sa fiksnim kašnjenjem server može da "
  "smanji batch sporog radnika na najmanju dozvoljenu vrednost, a runda će i dalje "
  "trajati isto. Eksperiment iz ovog odeljka zato koristi kašnjenje po uzorku, jer "
  "je to jedini oblik sporosti na koji balansiranje uopšte može da utiče.")

h2("4.12 Snimanje stanja i oporavak servera")
p("Priča o toleranciji otkaza do ovog mesta bila je jednostrana. Radnik može da "
  "padne i obuka se nastavlja, ali server drži jedinu kopiju parametara i njegov "
  "pad znači gubitak celog runa. Snimanje stanja pretvara ishod „server je pao, "
  "run je izgubljen“ u „server je pao, pokreni ga ponovo i izgubi najviše nekoliko "
  "rundi“.")
p("Upis je atomski, po postupku iz odeljka 3.9:")
code([
    "with open(tmp, \"wb\") as f:",
    "    np.savez(f, W=W, round=..., samples=...)",
    "    f.flush()",
    "    os.fsync(f.fileno())",
    "os.replace(tmp, path)",
])
p("Privremeni fajl nastaje u istom direktorijumu kao odredišni, jer preko granice "
  "fajl sistema preimenovanje gubi atomičnost. Poziv fsync ide nad istim "
  "deskriptorom nad kojim je obavljen upis; raniji pokušaj da se fajl ponovo "
  "otvori samo za čitanje pa nad njim pozove fsync radi na Linux-u, ali na "
  "Windows-u podiže grešku, jer sistemski poziv zahteva pravo upisa. Učitavanje "
  "oštećenog snimka tretira se kao da snimak ne postoji, a ne kao fatalna greška, "
  "jer nedovršen snimak od ranijeg pada ne sme da zaustavi run koji pokušava da se "
  "oporavi upravo od tog pada.")
p("Eksperiment je izveden tako što je server pokrenut sa snimanjem na svakih deset "
  "rundi, zatim nasilno ugašen čim je prvi snimak nastao, i potom ponovo pokrenut "
  "sa zahtevom za nastavak. Rezultat je vidljiv u bazi: prvi run pokriva runde od "
  "1 do 30 i nema upisano vreme završetka, jer je proces ubijen; drugi run pokriva "
  "runde od 31 do 1.565. Nastavak je počeo tačno tamo gde je snimak stao, a "
  "finalna tačnost drugog runa iznosi 0,8990, što je identično vrednosti koju "
  "postiže neprekinuti run sa dva radnika iz tabele 3. Prekid i oporavak, dakle, "
  "nisu uticali na ishod obuke.")

h2("4.13 Skladištenje metrika")
p("Sve izmerene vrednosti upisuju se u SQLite bazu. Izbor baze umesto tekstualnih "
  "datoteka nije proizvoljan i postaje očigledan čim se pojavi više od jednog "
  "runa. Poređenje runova je spajanje po zajedničkom ključu, a spajanje nad "
  "direktorijumom datoteka lako dovodi do nenamernog poređenja runova sa "
  "različitim parametrima. Osim toga, konfiguracija runa i metrike po rundi imaju "
  "različitu granularnost, pa bi jedna tabela morala ili da ponavlja konfiguraciju "
  "u svakom redu ili da je izostavi. Najzad, diskretni događaji poput registracije "
  "ili izbacivanja radnika nemaju svoju kolonu u tabeli koja opisuje rundu.")
p("Šema zato ima četiri tabele, po jednu za svaki nivo granularnosti, kako je "
  "prikazano u tabeli 8.")

table("Tabela 8. Šema baze metrika i trenutni broj redova.",
      ["Tabela", "Jedan red je", "Redova", "Ključne kolone"],
      [["runs", "jedno pokretanje eksperimenta", "13",
        "label, mode, n_workers, balance, total_bytes"],
       ["rounds", "run × runda", "15.684",
        "test_acc, wall_clock, barrier_wait, mean_staleness"],
       ["worker_rounds", "run × runda × radnik", "33.248",
        "n_samples, seconds, staleness"],
       ["events", "jedan diskretan događaj", "76",
        "kind (registered/evicted), detail"]],
      [1.5, 2.2, 0.9, 3.0])

p("Tabela worker_rounds je ono što omogućava sliku 7 i tabelu 7, jer se podatak o "
  "pojedinačnom radniku ne može smestiti u red koji opisuje celu rundu. Tabela "
  "events čuva razlog izbacivanja i time razdvaja dva mehanizma otkaza iz tabele "
  "5.")
p("Što se konkurentnosti tiče, u bazu piše isključivo server, pa postoji tačno "
  "jedan pisac i poznato ograničenje SQLite-a u pogledu više paralelnih pisaca ne "
  "dolazi do izražaja. Režim upisa unapred (WAL) ipak je uključen, da bi skripta "
  "za crtanje figura mogla da čita bazu dok run još traje; čitaoci se otvaraju "
  "isključivo za čitanje. Redovi sa metrikama potvrđuju se svakih pedeset rundi, "
  "kao i neposredno pre svakog snimanja stanja, čime se obezbeđuje da posle pada "
  "servera baza uvek sadrži bar one runde koje pokriva i poslednji snimak.")

h2("4.14 Ograničenja")
p("Rezultate treba čitati sa nekoliko ograda.")
bullets([
    "Svi čvorovi rade na jednoj mašini i komuniciraju preko lokalne mrežne petlje. "
    "Kašnjenje i propusni opseg time su znatno povoljniji nego u stvarnoj mreži, "
    "što posebno utiče na zaključak o kompresiji iz odeljka 4.10.",
    "Signal zatvorene veze dostupan je zato što otkaz pogađa proces na istoj "
    "mašini. Za udaljene otkaze i mrežne podele raspoloživ bi bio samo tajmaut.",
    "Model je linearan i mali. Odnos vremena računanja i vremena komunikacije kod "
    "dubokih mreža bitno je drugačiji, pa bi i optimalne odluke bile drugačije.",
    "Server je i dalje jedina tačka otkaza. Snimanje stanja skraćuje oporavak, ali "
    "ga ne uklanja; za to bi bila potrebna replikacija servera i izbor vođe.",
    "Broj čvorova u eksperimentima je najviše četiri. Efekti poput rasta "
    "zastarelosti sa veličinom klastera na toj skali se tek naziru.",
])

pagebreak()

# ========================================================== 5. ZAKLJUCAK =====
h1("5. Zaključak")
p("Implementiran je i izmeren distribuirani okvir za obuku modela mašinskog "
  "učenja po topologiji parameter servera, napisan direktno nad TCP soketima. "
  "Sistem podržava sinhronu i asinhronu obuku, detekciju otkaza na dva nezavisna "
  "signala, lanac interceptora kao middleware sloj, dinamičko balansiranje "
  "opterećenja i atomsko snimanje stanja sa nastavkom obuke.")
p("Merenja daju nekoliko zaključaka koji su konkretniji od očekivanih.")
p("Sinhrona distribuirana obuka reprodukuje sekvencijalno računanje. Prvi prolaz "
  "kroz skup podataka daje brojčano identične vrednosti funkcije gubitka, a "
  "kasnije razilaženje ostaje ispod 0,0055 apsolutne razlike u tačnosti i potiče "
  "od sastava batch-eva na granici epohe, a ne od sinhronizacije.")
p("Paralelizacija skalira propusnost, ali ne i vreme do zadate tačnosti. Četiri "
  "čvora obrađuju 3,52 puta više uzoraka u sekundi, a do iste tačnosti stižu za "
  "isto vreme kao jedan čvor, jer veći globalni batch pri nepromenjenom koraku "
  "učenja znači manje napretka po koraku. Skaliranje broja čvorova bez skaliranja "
  "koraka učenja daje sistem koji brže troši podatke bez ubrzanja obuke.")
p("Asinhroni režim ostvaruje ono što sinhroni ne može: pri istom broju koraka "
  "optimizacije i praktično istoj tačnosti, četiri čvora rade 2,62 puta brže od "
  "jednog. Cena je prosečna zastarelost gradijenata od 2,97 rundi, koja na ovoj "
  "veličini klastera nije umanjila tačnost.")
p("Dva mehanizma otkaza razlikuju se u ceni detekcije za oko 378 puta: 13,5 "
  "milisekundi kada se veza zatvori, prema 5,1 sekunde kada istekne tajmaut. Ta "
  "razlika je praktična posledica nemogućnosti da se u asinhronoj mreži spor čvor "
  "razlikuje od mrtvog, i ne može se ukloniti podešavanjem, nego samo zameniti "
  "drugim kompromisom.")
p("Kompresija nije bezuslovno korisna. Prelazak na jednostruku preciznost "
  "prepolovljava saobraćaj bez ikakve cene, ali dodatna deflate kompresija, iako "
  "smanjuje saobraćaj na 27,7 odsto polazne vrednosti, produžava izvršavanje 2,37 "
  "puta, jer je gust niz brojeva u pokretnom zarezu teško kompresibilan, a prenos "
  "preko lokalne petlje jeftin.")
p("Dinamičko balansiranje rešava problem sporog čvora u meri u kojoj je uzrok "
  "sporosti srazmeran količini posla. Čekanje na barijeri smanjeno je petnaest "
  "puta, a ukupno vreme 2,31 put, pri čemu izmerena raspodela batch-eva odgovara "
  "teorijskoj vrednosti izvedenoj iz odnosa brzina. Uslov ispravnosti je da prosek "
  "gradijenata bude težinski, jer bi obično usrednjavanje tiho promenilo funkciju "
  "koja se optimizuje.")
p("Zajednički zaključak svih merenja jeste da u distribuiranom sistemu ne postoji "
  "poboljšanje bez cene. Više čvorova donosi propusnost i oduzima efikasnost "
  "koraka optimizacije. Uklanjanje barijere donosi brzinu i oduzima "
  "determinisanost. Kraći tajmaut donosi brži oporavak i oduzima toleranciju prema "
  "sporim čvorovima. Kompresija donosi propusni opseg i oduzima procesorsko vreme. "
  "Vrednost izmerenog sistema nije u tome što je izabrao ispravnu stranu svakog od "
  "ovih kompromisa, nego u tome što ih čini vidljivim i merljivim.")

# ========================================================= 6. LITERATURA =====
h1("6. Literatura")

lit = [
    "Tanenbaum, A. S., Van Steen, M. (2017). Distributed Systems: Principles and "
    "Paradigms, 3. izdanje. Pearson Education. (Poglavlja o middleware-u, "
    "interceptorima i detekciji otkaza.)",
    "Coulouris, G., Dollimore, J., Kindberg, T., Blair, G. (2011). Distributed "
    "Systems: Concepts and Design, 5. izdanje. Addison-Wesley. (Modeli otkaza, "
    "sinhroni i asinhroni sistemi.)",
    "Li, M., Andersen, D. G., Park, J. W., Smola, A. J., Ahmed, A., Josifovski, "
    "V., Long, J., Shekita, E. J., Su, B.-Y. (2014). Scaling Distributed Machine "
    "Learning with the Parameter Server. Proceedings of the 11th USENIX Symposium "
    "on Operating Systems Design and Implementation (OSDI '14), str. 583–598.",
    "Dean, J., Corrado, G. S., Monga, R., Chen, K., Devin, M., Le, Q. V., Mao, M. "
    "Z., Ranzato, M., Senior, A., Tucker, P., Yang, K., Ng, A. Y. (2012). Large "
    "Scale Distributed Deep Networks. Advances in Neural Information Processing "
    "Systems 25 (NIPS 2012), str. 1223–1231.",
    "Recht, B., Re, C., Wright, S., Niu, F. (2011). Hogwild!: A Lock-Free Approach "
    "to Parallelizing Stochastic Gradient Descent. Advances in Neural Information "
    "Processing Systems 24 (NIPS 2011), str. 693–701.",
    "Chandra, T. D., Toueg, S. (1996). Unreliable Failure Detectors for Reliable "
    "Distributed Systems. Journal of the ACM, 43(2), str. 225–267.",
    "Dean, J., Ghemawat, S. (2004). MapReduce: Simplified Data Processing on Large "
    "Clusters. Proceedings of the 6th Symposium on Operating Systems Design and "
    "Implementation (OSDI '04), str. 137–150.",
    "Stevens, W. R., Fenner, B., Rudoff, A. M. (2003). UNIX Network Programming, "
    "Volume 1: The Sockets Networking API, 3. izdanje. Addison-Wesley. (Semantika "
    "poziva recv i uokvirivanje poruka nad tokom bajtova.)",
    "Deutsch, P. (1996). DEFLATE Compressed Data Format Specification version 1.3. "
    "RFC 1951, Internet Engineering Task Force. https://www.rfc-editor.org/rfc/rfc1951",
    "Harris, C. R. i saradnici (2020). Array programming with NumPy. Nature, 585, "
    "str. 357–362.",
    "LeCun, Y., Cortes, C., Burges, C. J. C. The MNIST Database of Handwritten "
    "Digits. http://yann.lecun.com/exdb/mnist/ (pristupljeno preko OpenML servisa).",
    "SQLite Consortium. Write-Ahead Logging. SQLite dokumentacija. "
    "https://www.sqlite.org/wal.html",
    "Python Software Foundation. Python 3 Standard Library dokumentacija: moduli "
    "socket, select, struct, zlib i sqlite3. https://docs.python.org/3/library/",
]

for i, ref in enumerate(lit, 1):
    lp = doc.add_paragraph()
    lp.paragraph_format.left_indent = Cm(0.9)
    lp.paragraph_format.first_line_indent = Cm(-0.9)
    lp.paragraph_format.space_after = Pt(6)
    lp.paragraph_format.line_spacing = 1.0
    lp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Tab, a ne razmaci: viseci uvlaka daje implicitni tab stop na 0,9 cm, pa
    # tekst svake stavke pocinje na istom mestu i kad je red poravnat obostrano.
    r = lp.add_run(f"[{i}]\t{ref}")
    r.font.name = FONT
    r.font.size = Pt(11)

# ------------------------------------------------------------- podnozje -----
footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
field(footer, "PAGE")
for r in footer.runs:
    r.font.name = FONT
    r.font.size = Pt(10)

doc.save(OUT)
print("Sacuvano:", OUT)
